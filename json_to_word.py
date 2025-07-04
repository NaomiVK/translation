import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class JSONToWordConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.doc.styles
        
        # Main title style
        if 'Main Title' not in [style.name for style in styles]:
            main_title = styles.add_style('Main Title', WD_STYLE_TYPE.PARAGRAPH)
            main_title.font.size = Inches(0.25)  # 18pt
            main_title.font.bold = True
            main_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            main_title.paragraph_format.space_after = Inches(0.2)
        
        # Heading 1 style
        if 'Custom Heading 1' not in [style.name for style in styles]:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.size = Inches(0.2)  # 14pt
            heading1.font.bold = True
            heading1.paragraph_format.space_before = Inches(0.15)
            heading1.paragraph_format.space_after = Inches(0.1)
        
        # Heading 2 style
        if 'Custom Heading 2' not in [style.name for style in styles]:
            heading2 = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.size = Inches(0.18)  # 13pt
            heading2.font.bold = True
            heading2.paragraph_format.space_before = Inches(0.12)
            heading2.paragraph_format.space_after = Inches(0.08)
    
    def add_breadcrumb(self, breadcrumb_data):
        """Add breadcrumb navigation"""
        breadcrumb_text = " > ".join([item['text'] for item in breadcrumb_data])
        p = self.doc.add_paragraph()
        p.add_run(breadcrumb_text).italic = True
        p.paragraph_format.space_after = Inches(0.15)
    
    def process_content_item(self, item):
        """Process a single content item based on its type"""
        item_type = item.get('type', '')
        
        if item_type == 'main_title':
            p = self.doc.add_paragraph(item['text'], style='Main Title')
            
        elif item_type == 'heading1':
            p = self.doc.add_paragraph(item['text'], style='Custom Heading 1')
            
        elif item_type == 'heading2':
            p = self.doc.add_paragraph(item['text'], style='Custom Heading 2')
            
        elif item_type == 'paragraph':
            p = self.doc.add_paragraph(item['text'])
            # Handle links if present
            if 'link' in item:
                p.add_run(" (")
                p.add_run(item['link']['text']).italic = True
                p.add_run(")")
                
        elif item_type == 'numbered_item':
            # Add the numbered item title
            p = self.doc.add_paragraph()
            p.add_run(f"{item['number']}. {item['title']}").bold = True
            
            # Add description if present
            if 'description' in item:
                desc_p = self.doc.add_paragraph(item['description'])
                desc_p.paragraph_format.left_indent = Inches(0.5)
            
            # Add additional info if present
            if 'additional_info' in item:
                info_p = self.doc.add_paragraph(item['additional_info'])
                info_p.paragraph_format.left_indent = Inches(0.5)
            
            # Handle subsections
            if 'subsections' in item:
                for subsection in item['subsections']:
                    self.process_subsection(subsection)
            
            # Handle options
            if 'options' in item:
                for option in item['options']:
                    self.process_option(option)
            
            # Handle special sections
            if 'special_section' in item:
                self.process_special_section(item['special_section'])
                
        elif item_type == 'feature_box':
            # Add a heading for features
            p = self.doc.add_paragraph("Features", style='Custom Heading 2')
            for feature in item['items']:
                feature_p = self.doc.add_paragraph()
                feature_p.add_run(f"• {feature['title']}").bold = True
                feature_p.add_run(f": {feature['description']}")
                feature_p.paragraph_format.left_indent = Inches(0.25)
    
    def process_subsection(self, subsection):
        """Process a subsection within a numbered item"""
        # Add subsection title
        title_p = self.doc.add_paragraph(subsection['title'])
        title_p.paragraph_format.left_indent = Inches(0.5)
        title_p.runs[0].bold = True
        
        # Add items
        if 'items' in subsection:
            for item in subsection['items']:
                if isinstance(item, dict):
                    item_p = self.doc.add_paragraph()
                    item_p.paragraph_format.left_indent = Inches(0.75)
                    
                    if 'date' in item:
                        item_p.add_run(f"• {item['date']}: ").bold = True
                        item_p.add_run(item['description'])
                        if 'note' in item:
                            item_p.add_run(f" ({item['note']})")
                    else:
                        item_p.add_run(f"• {item}")
                else:
                    item_p = self.doc.add_paragraph(f"• {item}")
                    item_p.paragraph_format.left_indent = Inches(0.75)
    
    def process_option(self, option):
        """Process options within a numbered item"""
        # Add option title
        title_p = self.doc.add_paragraph(option['title'])
        title_p.paragraph_format.left_indent = Inches(0.5)
        title_p.runs[0].bold = True
        
        # Add option items
        for item in option['items']:
            item_p = self.doc.add_paragraph(f"• {item}")
            item_p.paragraph_format.left_indent = Inches(0.75)
    
    def process_special_section(self, section):
        """Process special sections like info boxes"""
        # Add section title
        title_p = self.doc.add_paragraph(section['title'])
        title_p.paragraph_format.left_indent = Inches(0.5)
        title_p.runs[0].bold = True
        title_p.paragraph_format.space_before = Inches(0.1)
        
        # Add section content
        for content_item in section['content']:
            if content_item['type'] == 'paragraph':
                p = self.doc.add_paragraph(content_item['text'])
                p.paragraph_format.left_indent = Inches(0.75)
                
                # Handle links
                if 'link' in content_item:
                    p.add_run(" (")
                    p.add_run(content_item['link']['text']).italic = True
                    p.add_run(")")
                    
            elif content_item['type'] == 'note':
                p = self.doc.add_paragraph(content_item['text'])
                p.paragraph_format.left_indent = Inches(0.75)
                p.runs[0].italic = True
                
            elif content_item['type'] == 'emphasis':
                p = self.doc.add_paragraph(content_item['text'])
                p.paragraph_format.left_indent = Inches(0.75)
                p.runs[0].bold = True
                
            elif content_item['type'] == 'subsection':
                subsection_p = self.doc.add_paragraph(content_item['title'])
                subsection_p.paragraph_format.left_indent = Inches(0.75)
                subsection_p.runs[0].bold = True
                
                text_p = self.doc.add_paragraph(content_item['text'])
                text_p.paragraph_format.left_indent = Inches(1.0)
                
                for item in content_item['items']:
                    item_p = self.doc.add_paragraph(f"• {item}")
                    item_p.paragraph_format.left_indent = Inches(1.25)
    
    def convert_json_to_word(self, json_file):
        """Convert a JSON file to a Word document"""
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Add breadcrumb if present
        if 'breadcrumb' in data:
            self.add_breadcrumb(data['breadcrumb'])
        
        # Process content items
        for item in data['content']:
            self.process_content_item(item)
        
        # Generate output filename
        base_name = os.path.splitext(json_file)[0]
        output_file = f"{base_name}.docx"
        
        # Save the document
        self.doc.save(output_file)
        return output_file

def main():
    """Convert all translated JSON files to Word documents"""
    json_files = [
        'get_ready_c4ai_aya_expanse_32b_20250703.json',
        'get_ready_mistralai_mixtral_8x7b_instruct_20250703.json',
        'get_ready_openai_gpt_4o_mini_20250703.json'
    ]
    
    for json_file in json_files:
        if os.path.exists(json_file):
            print(f"Converting {json_file}...")
            converter = JSONToWordConverter()
            output_file = converter.convert_json_to_word(json_file)
            print(f"Created: {output_file}")
        else:
            print(f"File not found: {json_file}")

if __name__ == "__main__":
    main()